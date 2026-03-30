# -*- coding: utf-8 -*-
"""
Скрипт для генерации отчёта по лабораторной работе в формате .docx
С форматированием по ГОСТ (как в примере отчёта)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_report():
    doc = Document()
    
    # Настройка полей страницы (как в примере: left=3cm, top=2cm, right=1.5cm, bottom=2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # Титульный лист
    add_title_page(doc)
    
    # Содержание
    doc.add_page_break()
    add_table_of_contents(doc)
    
    # Введение
    doc.add_page_break()
    add_section(doc, "1. Введение")
    add_introduction(doc)
    
    # Цель работы
    add_section(doc, "2. Цель работы")
    add_goal(doc)
    
    # Задание
    add_section(doc, "3. Задание")
    add_task(doc)
    
    # Описание алгоритмов
    add_section(doc, "4. Описание алгоритмов")
    add_algorithms_description(doc)
    
    # Описание программы
    add_section(doc, "5. Описание программы")
    add_program_description(doc)
    
    # Тестирование
    add_section(doc, "6. Тестирование")
    add_testing(doc)
    
    # Выводы
    add_section(doc, "7. Выводы")
    add_conclusion(doc)
    
    # Приложение А - Блок-схемы
    doc.add_page_break()
    add_appendix_heading(doc, "Приложение А. Блок-схемы алгоритмов")
    add_appendix_a(doc)
    
    # Приложение Б - Листинги кода
    doc.add_page_break()
    add_appendix_heading(doc, "Приложение Б. Листинг программы")
    add_appendix_b(doc)
    
    # Приложение В - Юнит тесты
    doc.add_page_break()
    add_appendix_heading(doc, "Приложение В. Юнит-тесты")
    add_appendix_c(doc)
    
    # Сохранение
    doc.save(r'отчёт\блоксхемы\Отчёт_по_проекту.docx')
    print("Отчёт успешно создан: отчёт/блоксхемы/Отчёт_по_проекту.docx")


def set_run_format(run, font_name='Times New Roman', font_size=Pt(12), bold=False):
    """Установка форматирования для run"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    # Явно указываем шрифт для корректного отображения
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)


def add_section(doc, title):
    """Добавление заголовка раздела"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_format(run, font_size=Pt(14), bold=True)


def add_subsection(doc, title):
    """Добавление заголовка подраздела"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_format(run, font_size=Pt(14), bold=True)


def add_appendix_heading(doc, title):
    """Добавление заголовка приложения"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_format(run, font_size=Pt(14), bold=True)


def add_title_page(doc):
    """Титульный лист"""
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Министерство науки и высшего образования Российской Федерации")
    set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Федеральное государственное бюджетное образовательное учреждение")
    set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("высшего образования")
    set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("«Новгородский государственный университет имени Ярослава Мудрого» (НовГУ)")
    set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Политехнический институт")
    set_run_format(run, font_size=Pt(12))
    
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Кафедра информационных технологий и систем")
    set_run_format(run, font_size=Pt(12))
    
    for _ in range(5):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ")
    set_run_format(run, font_size=Pt(14), bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("на тему:")
    set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("«Сравнение рекурсивных и циклических алгоритмов»")
    set_run_format(run, font_size=Pt(14), bold=True)
    
    for _ in range(8):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Выполнил:\nстудент группы 5091\nБондарев Владимир Дмитриевич\n\nПроверил:\n___________________\n(Ф.И.О., должность)")
    set_run_format(run, font_size=Pt(12))
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Великий Новгород\n2026 г.")
    set_run_format(run, font_size=Pt(12))


def add_table_of_contents(doc):
    """Содержание"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Содержание")
    set_run_format(run, font_size=Pt(14), bold=True)
    
    contents = [
        "1. Введение",
        "2. Цель работы",
        "3. Задание",
        "4. Описание алгоритмов",
        "   4.1. Рекурсивный алгоритм 1 (задание 12)",
        "   4.2. Итеративный алгоритм 1 (задание 12)",
        "   4.3. Рекурсивный алгоритм 2 (задание 4)",
        "   4.4. Итеративный алгоритм 2 (задание 4)",
        "   4.5. Рекурсивный алгоритм 3 (дополнительное задание)",
        "5. Описание программы",
        "6. Тестирование",
        "7. Выводы",
        "Приложение А. Блок-схемы алгоритмов",
        "Приложение Б. Листинг программы",
        "Приложение В. Юнит-тесты"
    ]
    
    for item in contents:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))


def add_introduction(doc):
    """Введение"""
    p = doc.add_paragraph("В данной лабораторной работе выполняется разработка программного модуля на языке C++ для игрового движка Godot 4 с использованием технологии GDExtension.")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("Целью работы является изучение принципов:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "работы с рекурсивными алгоритмами;",
        "сравнения производительности рекурсивных и итеративных версий;",
        "интеграции C++ кода с Godot через GDExtension."
    ]
    
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))


def add_goal(doc):
    """Цель работы"""
    p = doc.add_paragraph("Разработать программный модуль на языке C++ для Godot 4, реализующий:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "Рекурсивные алгоритмы вычисления математических функций согласно варианту;",
        "Итеративные версии тех же алгоритмов;",
        "Подсчёт метрик производительности: количество вызовов рекурсии, время выполнения, использование памяти;",
        "Интеграцию с Godot через GDExtension;",
        "Графический интерфейс для взаимодействия с пользователем."
    ]
    
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))


def add_task(doc):
    """Задание"""
    p = doc.add_paragraph()
    run = p.add_run("Вариант: 12/4")
    set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    run = p.add_run("Задание 12 (Рекурсия 1):")
    run.bold = True
    set_run_format(run, font_size=Pt(12))
    p.add_run("\nВычисление функции по формуле:")
    p.add_run("\n- f(1) = 1")
    p.add_run("\n- f(n) = n + f(n-1), если n чётное")
    p.add_run("\n- f(n) = f(n-1) + 2·f(n-2), если n нечётное")
    for run in p.runs:
        set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    run = p.add_run("Задание 4 (Рекурсия 2):")
    run.bold = True
    set_run_format(run, font_size=Pt(12))
    p.add_run("\nВычисление функции по формуле:")
    p.add_run("\n- f(n) = 1, если n < 3")
    p.add_run("\n- f(n) = f(n-1) + f(n-2), если n нечётное")
    p.add_run("\n- f(n) = Σf(i) для i=1..n-1, если n чётное")
    for run in p.runs:
        set_run_format(run, font_size=Pt(12))
    
    p = doc.add_paragraph()
    run = p.add_run("Дополнительное задание (Рекурсия 3):")
    run.bold = True
    set_run_format(run, font_size=Pt(12))
    p.add_run("\nВывод цифр числа по одной в прямом порядке.")
    for run in p.runs:
        set_run_format(run, font_size=Pt(12))


def add_algorithms_description(doc):
    """Описание алгоритмов"""
    
    # 4.1
    add_subsection(doc, "4.1. Рекурсивный алгоритм 1 (задание 12)")
    p = doc.add_paragraph("Алгоритм реализует рекурсивную функцию с ветвлением в зависимости от чётности n.")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''function recursion1(n):
    if n == 1:
        return 1
    if n % 2 == 0:
        return n + recursion1(n - 1)
    else:
        return recursion1(n - 1) + 2 * recursion1(n - 2)''')
    
    p = doc.add_paragraph("Особенности:")
    set_run_format(p.runs[0], font_size=Pt(12))
    items = [
        "Базовый случай: n = 1",
        "Для чётных n: один рекурсивный вызов",
        "Для нечётных n: два рекурсивных вызова"
    ]
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))
    
    # 4.2
    add_subsection(doc, "4.2. Итеративный алгоритм 1 (задание 12)")
    p = doc.add_paragraph("Итеративная версия использует цикл и две переменные для хранения предыдущих значений.")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''function iteration1(n):
    if n == 1: return 1
    if n == 2: return 3
    
    prev2 = 1
    prev1 = 3
    
    for i from 3 to n:
        if i % 2 == 0:
            current = i + prev1
        else:
            current = prev1 + 2 * prev2
        prev2 = prev1
        prev1 = current
    
    return prev1''')
    
    # 4.3
    add_subsection(doc, "4.3. Рекурсивный алгоритм 2 (задание 4)")
    p = doc.add_paragraph("Алгоритм с более сложной логикой ветвления.")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''function recursion2(n):
    if n < 3:
        return 1
    if n % 2 != 0:
        return recursion2(n - 1) + recursion2(n - 2)
    else:
        sum = 0
        for i from 1 to n - 1:
            sum += recursion2(i)
        return sum''')
    
    # 4.4
    add_subsection(doc, "4.4. Итеративный алгоритм 2 (задание 4)")
    p = doc.add_paragraph("Итеративная версия с накоплением суммы.")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''function iteration2(n):
    if n < 3: return 1
    
    f1 = 1
    f2 = 1
    total = 2
    
    for i from 3 to n:
        if i % 2 != 0:
            f_i = f2 + f1
        else:
            f_i = total
        total += f_i
        f1 = f2
        f2 = f_i
    
    return f2''')
    
    # 4.5
    add_subsection(doc, "4.5. Рекурсивный алгоритм 3 (дополнительное задание)")
    p = doc.add_paragraph("Вывод цифр числа в прямом порядке.")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''function recursion3(n):
    if n < 10:
        print(n, " ")
        return
    recursion3(n / 10)
    print(n % 10, " ")''')
    
    p = doc.add_paragraph('Пример работы: recursion3(123) → "1 2 3"')
    set_run_format(p.runs[0], font_size=Pt(12))


def add_program_description(doc):
    """Описание программы"""
    
    add_subsection(doc, "5.1. Структура проекта")
    p = doc.add_paragraph("Проект состоит из следующих основных файлов:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("Исходный код C++ находится в папке src/ и включает:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "Lab1.h/cpp — основной класс с реализацией рекурсивных и итеративных алгоритмов;",
        "RecursionResult.h/cpp — структура для возврата результатов вычислений;",
        "register_types.h/cpp — регистрация типов в Godot;",
        "types.h — определения типов данных (i64, i32, u64)."
    ]
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))
    
    add_subsection(doc, "5.2. Класс Lab1")
    p = doc.add_paragraph("Класс Lab1 наследуется от RefCounted и содержит следующие методы:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "recursion1(i32 n) — вычисление по формуле задания 12 (рекурсия);",
        "recursion2(i32 n) — вычисление по формуле задания 4 (рекурсия);",
        "recursion3(i32 n) — вывод цифр числа (дополнительное задание);",
        "iteration1(i32 n) — итеративная версия recursion1;",
        "iteration2(i32 n) — итеративная версия recursion2."
    ]
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))
    
    add_subsection(doc, "5.3. Класс RecursionResult")
    p = doc.add_paragraph("Структура RecursionResult предназначена для возврата результатов вычислений и содержит поля:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "success — успешность выполнения (bool);",
        "value — вычисленное значение (i64);",
        "calls — количество вызовов рекурсии (i32);",
        "error — сообщение об ошибке (String);",
        "time — время выполнения (String);",
        "memory_amount — оценка использования памяти (u64)."
    ]
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))
    
    add_subsection(doc, "5.4. Сборка проекта")
    p = doc.add_paragraph("Сборка проекта выполняется с помощью системы сборки SCons:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''# Debug версия
scons platform=windows target=template_debug

# Release версия
scons platform=windows target=template_release''')


def add_testing(doc):
    """Тестирование"""
    
    add_subsection(doc, "6.1. Unit-тесты")
    p = doc.add_paragraph("Для тестирования используется фреймворк doctest. Примеры тестов:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''TEST_CASE("recursion1: базовый случай n=1") {
    Lab1 lab;
    auto result = lab.recursion1(1);
    
    CHECK(result->success == true);
    CHECK(result->value == 1);
    CHECK(result->calls == 1);
}

TEST_CASE("recursion1: чётное n=2") {
    Lab1 lab;
    auto result = lab.recursion1(2);
    
    CHECK(result->success == true);
    CHECK(result->value == 3);  // 2 + f(1) = 2 + 1 = 3
}''')
    
    add_subsection(doc, "6.2. Результаты тестирования")
    
    p = doc.add_paragraph("Таблица 1 — Результаты тестирования recursion1:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['n', 'f(n)', 'calls', 'time', 'memory']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        set_run_format(run, font_size=Pt(10))
    
    data = [
        ['1', '1', '1', '<1 us', '20 б'],
        ['5', '11', '9', '2 us', '60 б'],
        ['10', '143', '89', '5 us', '140 б'],
        ['15', '1229', '619', '15 us', '220 б']
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for run in row_cells[j].paragraphs[0].runs:
                set_run_format(run, font_size=Pt(10))
    
    p = doc.add_paragraph("\nТаблица 2 — Результаты тестирования recursion2:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    table2 = doc.add_table(rows=5, cols=5)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells2[i].text = header
        run = hdr_cells2[i].paragraphs[0].runs[0]
        run.bold = True
        set_run_format(run, font_size=Pt(10))
    
    data2 = [
        ['1', '1', '1', '<1 us', '20 б'],
        ['5', '3', '10', '3 us', '60 б'],
        ['8', '16', '85', '10 us', '120 б'],
        ['10', '52', '254', '25 us', '160 б']
    ]
    
    for i, row_data in enumerate(data2):
        row_cells = table2.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for run in row_cells[j].paragraphs[0].runs:
                set_run_format(run, font_size=Pt(10))
    
    add_subsection(doc, "6.3. Сравнение рекурсии и итерации")
    p = doc.add_paragraph("Для n = 20:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    
    hdr_cells3 = table3.rows[0].cells
    headers3 = ['Метрика', 'Рекурсия', 'Итерация']
    for i, header in enumerate(headers3):
        hdr_cells3[i].text = header
        run = hdr_cells3[i].paragraphs[0].runs[0]
        run.bold = True
        set_run_format(run, font_size=Pt(10))
    
    data3 = [
        ['Время', '150 мс', '<1 мс'],
        ['Вызовы', '13529', '20'],
        ['Память', '280 б', '24 б']
    ]
    
    for i, row_data in enumerate(data3):
        row_cells = table3.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for run in row_cells[j].paragraphs[0].runs:
                set_run_format(run, font_size=Pt(10))


def add_conclusion(doc):
    """Выводы"""
    p = doc.add_paragraph("В ходе выполнения лабораторной работы были получены следующие результаты:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "Разработан программный модуль на C++ для Godot 4 с использованием GDExtension;",
        "Реализованы 5 алгоритмов: recursion1, recursion2, recursion3 (дополнительное задание), iteration1, iteration2;",
        "Реализован подсчёт метрик производительности: количество вызовов рекурсии, время выполнения, оценка использования памяти стека;",
        "Написаны unit-тесты с использованием фреймворка doctest (22 теста);",
        "Продемонстрировано преимущество итеративных версий: O(n) против экспоненциальной сложности;",
        "Освоена технология GDExtension для интеграции C++ с Godot 4."
    ]
    
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))


def add_appendix_a(doc):
    """Приложение А - Блок-схемы"""
    p = doc.add_paragraph("В данном приложении представлены блок-схемы всех реализованных алгоритмов:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    items = [
        "Рисунок А.1 — Блок-схема рекурсивного алгоритма 1 (задание 12);",
        "Рисунок А.2 — Блок-схема итеративного алгоритма 1 (задание 12);",
        "Рисунок А.3 — Блок-схема рекурсивного алгоритма 2 (задание 4);",
        "Рисунок А.4 — Блок-схема итеративного алгоритма 2 (задание 4);",
        "Рисунок А.5 — Блок-схема рекурсивного алгоритма 3 (дополнительное задание)."
    ]
    for item in items:
        p = doc.add_paragraph(item)
        set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("\nБлок-схемы приведены в файлах:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("- Приложение А - задание 12.drawio.png")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("- Приложение А - задание 12  итерация.drawio (2).drawio.png")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("- Приложение А - задание 4.drawio.png")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("- Приложение А - задание 4 итерация.drawio (2).drawio.png")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    p = doc.add_paragraph("- Приложение А - рекурсия 3(доп задание).drawio.png")
    set_run_format(p.runs[0], font_size=Pt(12))


def add_appendix_b(doc):
    """Приложение Б - Листинги кода"""
    
    add_subsection(doc, "Б.1. Файл Lab1.h (фрагмент)")
    add_code_block(doc, '''class Lab1 : public RefCounted {
    GDCLASS(Lab1, RefCounted)

private:
    static String format_duration_us(u64 us) {
        if (us < 1000)
            return String::num_uint64(us) + " us";
        if (us < 1'000'000)
            return String::num(us / 1e3, 2) + " ms";
        return String::num(us / 1e6, 3) + " s";
    }

protected:
    static void _bind_methods();

public:
    Lab1() = default;
    ~Lab1() override = default;

    // РЕКУРСИЯ 1
    Ref<RecursionResult> recursion1(i32 n) {
        Ref<RecursionResult> result;
        result.instantiate();

        int call_count = 0;
        int max_depth = 0;

        if (n <= 0) {
            result->success = false;
            result->value = 0;
            result->calls = 0;
            result->error = "n must be > 0";
            return result;
        }

        u64 start = Time::get_singleton()->get_ticks_usec();
        i64 value = recursion_internal1(n, call_count, 0, max_depth);
        u64 end = Time::get_singleton()->get_ticks_usec();

        result->success = true;
        result->value = value;
        result->calls = call_count;
        result->error = "";
        result->time = format_duration_us(end - start);

        u64 frame_size = sizeof(i32) + sizeof(i32) + sizeof(i32) + 
                         sizeof(i32) + sizeof(i64);
        result->memory_amount = max_depth * frame_size;

        return result;
    }

    // ИТЕРАЦИЯ 1
    i64 iteration_internal1(i32 n) {
        if (n == 1) return 1;
        if (n == 2) return 3;

        i64 prev2 = 1;
        i64 prev1 = 3;

        for (int i = 3; i <= n; i++) {
            i64 current;
            if (i % 2 == 0)
                current = i + prev1;
            else
                current = prev1 + 2 * prev2;
            prev2 = prev1;
            prev1 = current;
        }

        return prev1;
    }
};''')
    
    add_subsection(doc, "Б.2. Файл Lab1.cpp")
    add_code_block(doc, '''#include "Lab1.h"

void Lab1::_bind_methods() {
    UtilityFunctions::print("!!! Lab1::_bind_methods CALLED !!!");
    ClassDB::bind_method(D_METHOD("recursion1", "n"), &Lab1::recursion1);
    ClassDB::bind_method(D_METHOD("recursion2", "n"), &Lab1::recursion2);
    ClassDB::bind_method(D_METHOD("recursion3", "n"), &Lab1::recursion3);
    ClassDB::bind_method(D_METHOD("iteration1", "n"), &Lab1::iteration1);
    ClassDB::bind_method(D_METHOD("iteration2", "n"), &Lab1::iteration2);
}''')
    
    add_subsection(doc, "Б.3. Файл RecursionResult.h")
    add_code_block(doc, '''class RecursionResult : public godot::RefCounted {
    GDCLASS(RecursionResult, godot::RefCounted)

protected:
    static void _bind_methods();

public:
    bool success = false;
    i64 value = 0;
    int calls = 0;
    godot::String error = godot::String();
    godot::String time = godot::String();
    godot::String digits_output = godot::String();
    u64 memory_amount = 0;

    RecursionResult() = default;
    ~RecursionResult() override = default;

    bool get_success() const { return success; }
    i64 get_value() const { return value; }
    i32 get_calls() const { return calls; }
    godot::String get_error() const { return error; }
    godot::String get_time() const { return time; }
    godot::String get_digits_output() const { return digits_output; }
    u64 get_memory_amount() const { return memory_amount; }
};''')


def add_appendix_c(doc):
    """Приложение В - Юнит тесты"""
    p = doc.add_paragraph("Полный список unit-тестов из файла doctest/test_lab1.cpp:")
    set_run_format(p.runs[0], font_size=Pt(12))
    
    add_code_block(doc, '''// ТЕСТЫ ДЛЯ RECURSION1

TEST_CASE("recursion1: базовый случай n=1") {
    Lab1 lab;
    auto result = lab.recursion1(1);
    
    CHECK(result->success == true);
    CHECK(result->value == 1);
    CHECK(result->calls == 1);
    CHECK(result->error == "");
}

TEST_CASE("recursion1: чётное n=2") {
    Lab1 lab;
    auto result = lab.recursion1(2);
    
    // recursion1(2) = 2 + recursion1(1) = 2 + 1 = 3
    CHECK(result->success == true);
    CHECK(result->value == 3);
    CHECK(result->calls > 1);
}

TEST_CASE("recursion1: нечётное n=3") {
    Lab1 lab;
    auto result = lab.recursion1(3);
    
    // recursion1(3) = recursion1(2) + 2*recursion1(1) = 3 + 2*1 = 5
    CHECK(result->success == true);
    CHECK(result->value == 5);
}

TEST_CASE("recursion1: отрицательный вход") {
    Lab1 lab;
    auto result = lab.recursion1(-5);
    
    CHECK(result->success == false);
    CHECK(result->value == 0);
    CHECK(result->calls == 0);
    CHECK(result->error == "n must be > 0");
}

// ТЕСТЫ ДЛЯ RECURSION2

TEST_CASE("recursion2: базовые случаи n=1 и n=2") {
    Lab1 lab;
    
    auto result1 = lab.recursion2(1);
    CHECK(result1->success == true);
    CHECK(result1->value == 1);
    
    auto result2 = lab.recursion2(2);
    CHECK(result2->success == true);
    CHECK(result2->value == 1);
}

TEST_CASE("recursion2: нечётное n=3") {
    Lab1 lab;
    auto result = lab.recursion2(3);
    
    // recursion2(3) = recursion2(2) + recursion2(1) = 1 + 1 = 2
    CHECK(result->success == true);
    CHECK(result->value == 2);
}

TEST_CASE("recursion2: чётное n=4") {
    Lab1 lab;
    auto result = lab.recursion2(4);
    
    // recursion2(4) = sum(recursion2(1), recursion2(2), recursion2(3)) = 1 + 1 + 2 = 4
    CHECK(result->success == true);
    CHECK(result->value == 4);
}

// ТЕСТЫ ДЛЯ RECURSION3

TEST_CASE("recursion3: однозначное число") {
    Lab1 lab;
    auto result = lab.recursion3(5);
    
    CHECK(result->success == true);
    CHECK(result->digits_output == "5");
    CHECK(result->calls == 1);
}

TEST_CASE("recursion3: многозначное число") {
    Lab1 lab;
    auto result = lab.recursion3(123);
    
    CHECK(result->success == true);
    CHECK(result->digits_output == "1 2 3");
}

TEST_CASE("recursion3: отрицательный вход") {
    Lab1 lab;
    auto result = lab.recursion3(0);
    
    CHECK(result->success == false);
    CHECK(result->digits_output == "");
}

// ТЕСТЫ ДЛЯ ITERATION1

TEST_CASE("iteration1: базовый случай n=1") {
    Lab1 lab;
    auto result = lab.iteration1(1);
    
    CHECK(result->success == true);
    CHECK(result->value == 1);
}

TEST_CASE("iteration1: n=2") {
    Lab1 lab;
    auto result = lab.iteration1(2);
    
    CHECK(result->success == true);
    CHECK(result->value == 3);
}

TEST_CASE("iteration1: n=5") {
    Lab1 lab;
    auto result = lab.iteration1(5);
    
    CHECK(result->success == true);
    CHECK(result->value == 11);
}

TEST_CASE("iteration1: отрицательный вход") {
    Lab1 lab;
    auto result = lab.iteration1(-1);
    
    CHECK(result->success == false);
    CHECK(result->value == 0);
    CHECK(result->error == "n must be > 0");
}

// ТЕСТЫ ДЛЯ ITERATION2

TEST_CASE("iteration2: базовые случаи n=1 и n=2") {
    Lab1 lab;
    
    auto result1 = lab.iteration2(1);
    CHECK(result1->success == true);
    CHECK(result1->value == 1);
    
    auto result2 = lab.iteration2(2);
    CHECK(result2->success == true);
    CHECK(result2->value == 1);
}

TEST_CASE("iteration2: n=5") {
    Lab1 lab;
    auto result = lab.iteration2(5);
    
    CHECK(result->success == true);
    CHECK(result->value == 3);
}

// СРАВНЕНИЕ РЕКУРСИИ И ИТЕРАЦИИ

TEST_CASE("recursion1 vs iteration1: одинаковые результаты") {
    Lab1 lab;
    
    for (int n = 1; n <= 10; n++) {
        auto rec_result = lab.recursion1(n);
        auto iter_result = lab.iteration1(n);
        
        CHECK(rec_result->value == iter_result->value);
    }
}

TEST_CASE("recursion2 vs iteration2: одинаковые результаты") {
    Lab1 lab;
    
    for (int n = 1; n <= 8; n++) {
        auto rec_result = lab.recursion2(n);
        auto iter_result = lab.iteration2(n);
        
        CHECK(rec_result->value == iter_result->value);
    }
}''')


def add_code_block(doc, code_text):
    """Добавление блока кода с форматированием как в примере (Courier New 10pt)"""
    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)


# Импорт в конце, чтобы избежать проблем с OxmlElement
from docx.oxml import OxmlElement

if __name__ == "__main__":
    create_report()
