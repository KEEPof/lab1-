# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm

doc = Document(r'C:\Users\Кип\godot\projects\labs\godot-cpp-template\отчёт\блоксхемы\пример отчёта.docx')

print('=== ДЕТАЛЬНЫЙ АНАЛИЗ ФОРМАТИРОВАНИЯ ===\n')

# Анализ заголовков
print('--- ЗАГОЛОВОК 1 (Heading 1) ---')
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Heading 1':
        pf = p.paragraph_format
        print(f'Параграф {i}: "{p.text[:50]}"')
        print(f'  space_before: {pf.space_before} ({pf.space_before.pt if pf.space_before else None} pt)')
        print(f'  space_after: {pf.space_after} ({pf.space_after.pt if pf.space_after else None} pt)')
        print(f'  line_spacing: {pf.line_spacing}')
        print(f'  alignment: {pf.alignment}')
        if p.runs:
            run = p.runs[0]
            print(f'  font: {run.font.name}, size: {run.font.size}, bold: {run.font.bold}')
        break

print('\n--- NORMAL (основной текст) ---')
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Normal' and p.text.strip():
        pf = p.paragraph_format
        print(f'Параграф {i}: "{p.text[:50]}"')
        print(f'  space_before: {pf.space_before} ({pf.space_before.pt if pf.space_before else None} pt)')
        print(f'  space_after: {pf.space_after} ({pf.space_after.pt if pf.space_after else None} pt)')
        print(f'  line_spacing: {pf.line_spacing}')
        print(f'  first_line_indent: {pf.first_line_indent}')
        if p.runs:
            run = p.runs[0]
            print(f'  font: {run.font.name}, size: {run.font.size}')
        if i > 45:
            break

print('\n--- ТАБЛИЦЫ ---')
for i, table in enumerate(doc.tables):
    print(f'\nТаблица {i}: {len(table.rows)} rows x {len(table.columns)} cols')
    # Первая ячейка
    cell = table.cell(0, 0)
    print(f'  Первая ячейка: "{cell.text[:50]}"')
    for p in cell.paragraphs:
        pf = p.paragraph_format
        print(f'    space_before: {pf.space_before}, space_after: {pf.space_after}')
        if p.runs:
            print(f'    font: {p.runs[0].font.name if p.runs else None}, size: {p.runs[0].font.size if p.runs else None}')

print('\n--- ЛИСТИНГИ КОДА ---')
# Ищем листинги - обычно это Normal с моноширинным шрифтом
for i, p in enumerate(doc.paragraphs):
    if p.runs and len(p.runs) > 0:
        run = p.runs[0]
        if run.font.name and ('Courier' in run.font.name or 'Consolas' in run.font.name or 'Lucida' in run.font.name):
            print(f'Параграф {i}: font={run.font.name}, size={run.font.size}')
            print(f'  Текст: {p.text[:60]}')
