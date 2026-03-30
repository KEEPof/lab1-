# -*- coding: utf-8 -*-
from docx import Document
import os

path = r'C:\Users\Кип\godot\projects\labs\godot-cpp-template\отчёт\блоксхемы\пример отчёта.docx'
print(f'File exists: {os.path.exists(path)}')

doc = Document(path)
print(f'Paragraphs count: {len(doc.paragraphs)}')
print(f'Sections count: {len(doc.sections)}')

print('\n=== СТРУКТУРА ДОКУМЕНТА ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text[:80] if p.text else '[EMPTY]'
    style = p.style.name if p.style else 'None'
    print(f'{i}: [{style}] {text}')

print('\n=== СЕКЦИИ ===')
for i, section in enumerate(doc.sections):
    print(f'Section {i}:')
    print(f'  top_margin: {section.top_margin}')
    print(f'  bottom_margin: {section.bottom_margin}')
    print(f'  left_margin: {section.left_margin}')
    print(f'  right_margin: {section.right_margin}')

print('\n=== СТИЛИ ===')
for style in doc.styles:
    print(f'  {style.name}')
