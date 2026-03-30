# -*- coding: utf-8 -*-
from docx import Document

example = Document(r'отчёт\блоксхемы\пример отчёта.docx')
new_report = Document(r'отчёт\блоксхемы\Отчёт_по_проекту_итог.docx')

print("=== ИТОГОВОЕ СРАВНЕНИЕ ===\n")

print("--- ЗАГОЛОВОК 1 (РАЗДЕЛ) ---")
print("Пример:")
for p in example.paragraphs:
    if p.style and p.style.name == 'Heading 1' and p.text.strip():
        run = p.runs[0] if p.runs else None
        print(f'  font: {run.font.name if run and run.font else None}, size: {run.font.size.pt if run and run.font.size else None} pt, bold: {run.font.bold if run else None}')
        print(f'  Текст: {p.text[:40]}')
        break

print("\nНовый отчёт:")
for p in new_report.paragraphs:
    if p.runs and p.runs[0].font.name == 'Times New Roman' and p.runs[0].font.size.pt == 14 and p.runs[0].font.bold:
        print(f'  font: {p.runs[0].font.name}, size: {p.runs[0].font.size.pt} pt, bold: {p.runs[0].font.bold}')
        print(f'  Текст: {p.text[:40]}')
        break

print("\n--- ОСНОВНОЙ ТЕКСТ ---")
print("Пример:")
for p in example.paragraphs:
    if p.style and p.style.name == 'Normal' and p.text.strip() and not p.runs[0].font.bold if p.runs else False:
        run = p.runs[0] if p.runs else None
        if run and run.font.size:
            print(f'  font: {run.font.name if run.font else None}, size: {run.font.size.pt if run.font.size else None} pt')
            print(f'  Текст: {p.text[:40]}')
            break

print("\nНовый отчёт:")
for p in new_report.paragraphs:
    if p.runs and p.runs[0].font.name == 'Times New Roman' and p.runs[0].font.size.pt == 12:
        print(f'  font: {p.runs[0].font.name}, size: {p.runs[0].font.size.pt} pt')
        print(f'  Текст: {p.text[:40]}')
        break

print("\n--- КОД (Courier New) ---")
print("Пример:")
for p in example.paragraphs:
    if p.runs and p.runs[0].font.name == 'Courier New':
        print(f'  font: {p.runs[0].font.name}, size: {p.runs[0].font.size.pt} pt')
        print(f'  Текст: {p.text[:40]}')
        break

print("\nНовый отчёт:")
for p in new_report.paragraphs:
    if p.runs and p.runs[0].font.name == 'Courier New':
        print(f'  font: {p.runs[0].font.name}, size: {p.runs[0].font.size.pt} pt')
        print(f'  Текст: {p.text[:40]}')
        break

print("\n--- ТАБЛИЦЫ ---")
print(f'Пример: {len(example.tables)} таблиц')
print(f'Новый отчёт: {len(new_report.tables)} таблиц')

print("\n--- СТРУКТУРА НОВОГО ОТЧЁТА (первые 60 параграфов) ---")
for i, p in enumerate(new_report.paragraphs[:60]):
    style = p.style.name if p.style else 'Custom'
    text = p.text[:50] if p.text else '[EMPTY]'
    if p.runs:
        run = p.runs[0]
        font_info = f"{run.font.name}, {run.font.size.pt if run.font.size else '?'}pt"
        if run.font.bold: font_info += " bold"
    else:
        font_info = 'No runs'
    print(f'{i}: [{style}] {font_info} | {text}')
