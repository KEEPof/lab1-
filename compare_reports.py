# -*- coding: utf-8 -*-
from docx import Document

# Сравнение примера и нового отчёта
example = Document(r'отчёт\блоксхемы\пример отчёта.docx')
new_report = Document(r'отчёт\блоксхемы\Отчёт_по_проекту_новый.docx')

print("=== СРАВНЕНИЕ ФОРМАТИРОВАНИЯ ===\n")

print("--- ПРИМЕР (Heading 1) ---")
for p in example.paragraphs:
    if p.style and p.style.name == 'Heading 1':
        pf = p.paragraph_format
        run = p.runs[0] if p.runs else None
        print(f'  font: {run.font.name if run else None}')
        print(f'  size: {run.font.size.pt if run and run.font.size else None} pt')
        print(f'  bold: {run.font.bold if run else None}')
        break

print("\n--- НОВЫЙ ОТЧЁТ (Heading 1) ---")
for p in new_report.paragraphs:
    if p.style and p.style.name == 'Heading 1':
        pf = p.paragraph_format
        run = p.runs[0] if p.runs else None
        print(f'  font: {run.font.name if run else None}')
        print(f'  size: {run.font.size.pt if run and run.font.size else None} pt')
        print(f'  bold: {run.font.bold if run else None}')
        break

print("\n--- ПРИМЕР (Normal) ---")
for p in example.paragraphs:
    if p.style and p.style.name == 'Normal' and p.text.strip():
        run = p.runs[0] if p.runs else None
        print(f'  font: {run.font.name if run and run.font else None}')
        print(f'  size: {run.font.size.pt if run and run.font.size else None} pt')
        break

print("\n--- НОВЫЙ ОТЧЁТ (Normal) ---")
for p in new_report.paragraphs:
    if p.style and p.style.name == 'Normal' and p.text.strip():
        run = p.runs[0] if p.runs else None
        print(f'  font: {run.font.name if run and run.font else None}')
        print(f'  size: {run.font.size.pt if run and run.font.size else None} pt')
        break

print("\n--- ПРИМЕР (код - Courier) ---")
for p in example.paragraphs:
    if p.runs and p.runs[0].font.name and 'Courier' in p.runs[0].font.name:
        run = p.runs[0]
        print(f'  font: {run.font.name}')
        print(f'  size: {run.font.size.pt if run.font.size else None} pt')
        break

print("\n--- НОВЫЙ ОТЧЁТ (код - Courier) ---")
for p in new_report.paragraphs:
    if p.runs and p.runs[0].font.name and 'Courier' in p.runs[0].font.name:
        run = p.runs[0]
        print(f'  font: {run.font.name}')
        print(f'  size: {run.font.size.pt if run.font.size else None} pt')
        break

print("\n--- ТАБЛИЦЫ ---")
print(f'Пример таблиц: {len(example.tables)}')
print(f'Новый отчёт таблиц: {len(new_report.tables)}')

for i, table in enumerate(new_report.tables[:3]):
    print(f'\nТаблица {i}: {len(table.rows)} rows x {len(table.columns)} cols')
    cell = table.cell(0, 0)
    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        print(f'  Шрифт в ячейке: {run.font.name}, size: {run.font.size.pt if run.font.size else None} pt')

print("\n--- СТРУКТУРА НОВОГО ОТЧЁТА ---")
for i, p in enumerate(new_report.paragraphs[:50]):
    style = p.style.name if p.style else 'None'
    text = p.text[:60] if p.text else '[EMPTY]'
    print(f'{i}: [{style}] {text}')
